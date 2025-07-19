import os
import io
from docx import Document
from docx.shared import Pt
from django.conf import settings
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def fill_template_with_data(data):
    template_path = os.path.join(settings.BASE_DIR, 'resume_template_1.docx')
    doc = Document(template_path)

    full_name = f"{data['first_name']} {data['last_name']}".upper()
    phone = data.get('phone', '')
    email = data.get('email', '')
    linkedin = data.get('linkedin_profile', '')
    github = data.get('github', '')
    summary = data.get('summary', '')
    coursework = data.get('coursework', '')
    skills_list = data.get('skills', [])
    experiences = data.get('experiences', [])

    def format_paragraph(paragraph, text, bold=False, size=12, spacing=1.0):
        paragraph.clear()
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        paragraph.paragraph_format.line_spacing = spacing

    def add_horizontal_line(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        border.append(bottom)
        pPr.append(border)

    def is_empty(value):
        return not value or (isinstance(value, list) and all(not str(item).strip() for item in value))

    

    def format_cert(cert):
        return f"{cert['title']} - {cert['issuing_organization']}\n{cert.get('credential_url', '')}"

    def format_achievement(ach):
        return f"{ach['title']}: {ach['description']}"

    def insert_education_section(doc, paragraph, educations):
        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)
        parent.remove(paragraph._element)

        for edu in educations:
            # First line: Institution (bold) + Location (right-aligned)
            table1 = doc.add_table(rows=1, cols=2)
            table1.autofit = True

            inst_p = table1.rows[0].cells[0].paragraphs[0]
            inst_run = inst_p.add_run(edu['institution'])
            inst_run.bold = True
            inst_run.font.name = 'Times New Roman'
            inst_run.font.size = Pt(12)
            inst_p.paragraph_format.line_spacing = 1.0
            inst_p.paragraph_format.space_after = Pt(0)

            loc_p = table1.rows[0].cells[1].paragraphs[0]
            loc_p.alignment = 2  # right align
            loc_run = loc_p.add_run(edu['location'])
            loc_run.font.name = 'Times New Roman'
            loc_run.font.size = Pt(12)
            loc_p.paragraph_format.line_spacing = 1.0
            loc_p.paragraph_format.space_after = Pt(0)

            parent.insert(index, table1._element)
            index += 1

            # Second line: Degree (bold) + Dates (right-aligned)
            table2 = doc.add_table(rows=1, cols=2)
            table2.autofit = True

            deg_p = table2.rows[0].cells[0].paragraphs[0]
            deg_run = deg_p.add_run(edu['degree'])
            deg_run.font.name = 'Times New Roman'
            deg_run.font.size = Pt(12)
            deg_p.paragraph_format.line_spacing = 1.0
            deg_p.paragraph_format.space_after = Pt(0)

            date_text = f"{edu['start_date']} – {edu.get('end_date', 'Present')}"
            date_p = table2.rows[0].cells[1].paragraphs[0]
            date_p.alignment = 2  # right align
            date_run = date_p.add_run(date_text)
            date_run.font.name = 'Times New Roman'
            date_run.font.size = Pt(12)
            date_p.paragraph_format.line_spacing = 1.0
            date_p.paragraph_format.space_after = Pt(0)

            parent.insert(index, table2._element)
            index += 1

            # Description as bullet points
            desc_points = [pt.strip() for pt in edu['description'].split('.') if pt.strip()]
            for line in desc_points:
                p = doc.add_paragraph(f"\u2022 {line}")
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.first_line_indent = Pt(-6)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(4)
                run = p.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                parent.insert(index, p._element)
                index += 1

    def insert_experience_table(doc, paragraph, experiences):
        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)
        parent.remove(paragraph._element)

        for exp in experiences:
            table = doc.add_table(rows=2, cols=2)
            table.autofit = True

            row1 = table.rows[0]
            company_p = row1.cells[0].paragraphs[0]
            run = company_p.add_run(exp['company_name'])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            company_p.paragraph_format.line_spacing = 1.0
            company_p.paragraph_format.space_after = Pt(0)

            location_p = row1.cells[1].paragraphs[0]
            location_p.alignment = 2
            run2 = location_p.add_run(exp['location'])
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            location_p.paragraph_format.line_spacing = 1.0
            location_p.paragraph_format.space_after = Pt(0)

            row2 = table.rows[1]
            role_p = row2.cells[0].paragraphs[0]
            run = role_p.add_run(exp['job_title'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            role_p.paragraph_format.line_spacing = 1.0
            role_p.paragraph_format.space_after = Pt(2)

            date_p = row2.cells[1].paragraphs[0]
            date_p.alignment = 2
            run2 = date_p.add_run(f"{exp['start_date']} – {exp.get('end_date', 'Present')}")
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            date_p.paragraph_format.line_spacing = 1.0
            date_p.paragraph_format.space_after = Pt(2)

            parent.insert(index, table._element)
            index += 1

            bullet_points = [point.strip() for point in exp['description'].split('.') if point.strip()]
            for line in bullet_points:
                p = doc.add_paragraph(f"\u2022 {line}")
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.first_line_indent = Pt(-6)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(6)
                run = p.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                parent.insert(index, p._element)
                index += 1

    def insert_project_section(doc, paragraph, projects):
        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)
        parent.remove(paragraph._element)

        for proj in projects:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True

            row = table.rows[0]
            title_p = row.cells[0].paragraphs[0]
            run = title_p.add_run(proj['title'])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            title_p.paragraph_format.line_spacing = 1.0
            title_p.paragraph_format.space_after = Pt(0)

            date_p = row.cells[1].paragraphs[0]
            date_p.alignment = 2
            run2 = date_p.add_run(f"{proj['start_date']} – {proj.get('end_date', 'Present')}")
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            date_p.paragraph_format.line_spacing = 1.0
            date_p.paragraph_format.space_after = Pt(0)

            parent.insert(index, table._element)
            index += 1

            bullet_points = [point.strip() for point in proj['description'].split('.') if point.strip()]
            for line in bullet_points:
                p = doc.add_paragraph(f"\u2022 {line}")
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.first_line_indent = Pt(-6)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(4)
                run = p.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                parent.insert(index, p._element)
                index += 1

    # Static text replacements
    replacements = {
        "{{ FULL_NAME }}": full_name,
        "{{ PHONE }}": phone,
        "{{ EMAIL }}": email,
        "{{ LINKEDIN }}": linkedin,
        "{{ GITHUB }}": github,
        "{{ SUMMARY }}": summary,
        "{{ COURSEWORK }}": coursework,
        "{{ CERTIFICATION_DETAILS }}": [format_cert(c) for c in data.get('certifications', [])],
        "{{ ACHIEVEMENTS }}": [format_achievement(a) for a in data.get('achievements', [])],
    }


    for i, paragraph in enumerate(doc.paragraphs):
        if "{{ EDUCATION_DETAILS }}" in paragraph.text:
            educations = data.get('educations', [])
            if is_empty(educations):
                if i > 0 and doc.paragraphs[i - 1]._element.getparent() is not None:
                    doc.paragraphs[i - 1]._element.getparent().remove(doc.paragraphs[i - 1]._element)
                if paragraph._element.getparent() is not None:
                    paragraph._element.getparent().remove(paragraph._element)
            else:
                insert_education_section(doc, paragraph, educations)
            break


    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            if isinstance(val, list):
                continue
            if key in paragraph.text:
                text = paragraph.text.replace(key, val)
                format_paragraph(paragraph, text, bold=(key == "{{ FULL_NAME }}"), size=(14 if key == "{{ FULL_NAME }}" else 12))

    for i, paragraph in enumerate(doc.paragraphs):
        for key, val in replacements.items():
            if isinstance(val, list) and key in paragraph.text:
                if is_empty(val):
                    if i > 0 and doc.paragraphs[i - 1]._element.getparent() is not None:
                        doc.paragraphs[i - 1]._element.getparent().remove(doc.paragraphs[i - 1]._element)
                    if paragraph._element.getparent() is not None:
                        paragraph._element.getparent().remove(paragraph._element)
                else:
                    parent = paragraph._element.getparent()
                    index = parent.index(paragraph._element)
                    parent.remove(paragraph._element)
                    for item in val:
                        p = doc.add_paragraph()
                        format_paragraph(p, item, spacing=1.0)
                        parent.insert(index, p._element)
                        index += 1
                break

    for i, paragraph in enumerate(doc.paragraphs):
        if "{{ EXPERIENCE_DETAILS }}" in paragraph.text:
            if is_empty(experiences):
                if i > 0 and doc.paragraphs[i - 1]._element.getparent() is not None:
                    doc.paragraphs[i - 1]._element.getparent().remove(doc.paragraphs[i - 1]._element)
                if paragraph._element.getparent() is not None:
                    paragraph._element.getparent().remove(paragraph._element)
            else:
                insert_experience_table(doc, paragraph, experiences)
            break

    for i, paragraph in enumerate(doc.paragraphs):
        if "{{ PROJECT_DETAILS }}" in paragraph.text:
            projects = data.get('projects', [])
            if is_empty(projects):
                if i > 0 and doc.paragraphs[i - 1]._element.getparent() is not None:
                    doc.paragraphs[i - 1]._element.getparent().remove(doc.paragraphs[i - 1]._element)
                if paragraph._element.getparent() is not None:
                    paragraph._element.getparent().remove(paragraph._element)
            else:
                insert_project_section(doc, paragraph, projects)
            break

    for paragraph in doc.paragraphs:
        if "{{ SKILLS }}" in paragraph.text:
            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)
            parent.remove(paragraph._element)
            if skills_list:
                num_cols = 4
                num_rows = (len(skills_list) + num_cols - 1) // num_cols
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.autofit = True
                idx = 0
                for row in table.rows:
                    for cell in row.cells:
                        if idx < len(skills_list):
                            p = cell.paragraphs[0]
                            run = p.add_run("\u2022 " + skills_list[idx])
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            p.paragraph_format.line_spacing = 0.8
                            p.paragraph_format.left_indent = Pt(12)
                            p.paragraph_format.first_line_indent = Pt(-6)
                            idx += 1
                        else:
                            cell.text = ""
                parent.insert(index, table._element)
            else:
                if index > 0 and doc.paragraphs[index - 1]._element.getparent() is not None:
                    doc.paragraphs[index - 1]._element.getparent().remove(doc.paragraphs[index - 1]._element)
            break

    for section in data.get('additional_sections', []):
        title = section.get('section_title', '').strip()
        content = section.get('content', '').strip()
        if title and content:
            p_title = doc.add_paragraph()
            run = p_title.add_run(title)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = True
            run.font.small_caps = True
            p_title.paragraph_format.line_spacing = 1.0
            add_horizontal_line(p_title)
            for line in content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(f"\u2022 {line.strip()}")
                    p.paragraph_format.line_spacing = 0.9
                    p.paragraph_format.left_indent = Pt(12)
                    p.paragraph_format.first_line_indent = Pt(-6)
                    run = p.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer